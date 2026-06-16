"""Generate the two hand-off documents as polished Word (.docx) files."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
NAVY = RGBColor(0x10, 0x18, 0x20)
ACCENT = RGBColor(0x0B, 0x5F, 0xFF)
GREY = RGBColor(0x66, 0x70, 0x7D)


def base_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for lvl, sz in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.color.rgb = NAVY
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.9)
    section.left_margin = section.right_margin = Inches(1)
    return doc


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    r = p.add_run("SKYRYSE")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    h = doc.add_paragraph()
    rh = h.add_run(title)
    rh.bold = True
    rh.font.size = Pt(19)
    rh.font.color.rgb = NAVY
    if subtitle:
        s = doc.add_paragraph()
        rs = s.add_run(subtitle)
        rs.italic = True
        rs.font.size = Pt(10.5)
        rs.font.color.rgb = GREY
    _hrule(doc.add_paragraph())


def _hrule(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B5FFF")
    pbdr.append(bottom)
    pPr.append(pbdr)


def meta_block(doc, rows):
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rl = p.add_run(f"{label}  ")
        rl.bold = True
        rl.font.size = Pt(10)
        rv = p.add_run(value)
        rv.font.size = Pt(10)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_table(doc, header, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(hdr[i], "10182 0".replace(" ", ""))
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(text))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            p.add_run(" " + it[1])
        else:
            p.add_run(it)


# ---------------------------------------------------------------------------
# Document 1 — System Requirements
# ---------------------------------------------------------------------------

def build_system_requirements():
    doc = base_doc()
    add_title(doc, "FAIR Review Tool — System Requirements",
              "Multi-User AWS Deployment")
    meta_block(doc, [
        ("Prepared for:", "Skyryse IT / Cloud Engineering"),
        ("Prepared by:", "Abel Napoleon, Supplier Quality"),
        ("Date:", "June 12, 2026"),
        ("Purpose:", "Move the FAIR Review pilot from a single workstation to a shared AWS-hosted service."),
    ])

    doc.add_heading("1. Scope & Expected Load", level=1)
    add_table(doc, ["Dimension", "Value"], [
        ["Concurrent users", "3–4 Supplier Quality engineers"],
        ["Reviews (volume)", "~40 FAIR packages (small)"],
        ["Typical package size", "~30 MB (PDFs, scanned certs, Excel FAIR forms)"],
        ["Usage pattern", "Occasional, interactive — not high-throughput batch"],
    ], widths=[2.2, 4.6])
    p = doc.add_paragraph()
    p.add_run("This is a low-volume internal tool. ").bold = True
    p.add_run("Size it for reliability and security, not scale. Do not over-provision.")

    doc.add_heading("2. Controlling Constraint — Export Control (read first)", level=1)
    doc.add_paragraph(
        "The data includes EAR / ECCN 7E994 controlled technical data (e.g., the 1101021 drawing). "
        "This drives the architecture:")
    bullets(doc, [
        ("Deploy in AWS GovCloud (US),", "not commercial AWS — GovCloud restricts operations to vetted U.S. persons and is the standard boundary for export-controlled workloads."),
        ("Run the AI review via Amazon Bedrock (Claude) inside GovCloud", "instead of the public Anthropic API, so document contents stay inside the AWS controlled boundary. (The pilot currently calls the public Anthropic API; this is a deliberate change for the compliant deployment.)"),
        ("Keep the service internal-only", "— reachable over the corporate network / VPN, not exposed to the public internet."),
    ])
    doc.add_paragraph(
        "Final confirmation of the above belongs to Skyryse's export-control / compliance officer "
        "(see the Compliance Summary document).")

    doc.add_heading("3. Target Architecture", level=1)
    doc.add_paragraph(
        "Engineers reach an internal Application Load Balancer (TLS) over VPN, which fronts the "
        "FastAPI application running on ECS Fargate (or EC2). The application connects to four AWS services:")
    add_table(doc, ["Service", "Role"], [
        ["Aurora PostgreSQL", "Structured data — part metadata, review results, audit log"],
        ["S3 bucket", "Binary documents (uploaded packages) and generated reports"],
        ["Amazon Bedrock", "Claude inference (GovCloud) — keeps data in the AWS boundary"],
        ["Secrets Manager", "Database credentials and keys"],
    ], widths=[2.2, 4.6])
    p = doc.add_paragraph()
    p.add_run("Storage split: ").bold = True
    p.add_run("Aurora holds structured data; S3 holds the binary documents and reports. "
              "Do not store large PDFs or images in the database.")

    doc.add_heading("4. Component Requirements", level=1)

    doc.add_heading("4.1  Compute (application server)", level=2)
    bullets(doc, [
        "ECS Fargate task: 2 vCPU / 4 GB RAM (or EC2 t3.medium). One task suffices; run 2 for high availability.",
        "Brief CPU spikes occur when rendering scanned PDF pages to images during conversion; 2 vCPU / 4 GB absorbs this for 3–4 concurrent users.",
        "Stateless — all persistent data lives in Aurora/S3, so tasks can be replaced freely.",
    ])

    doc.add_heading("4.2  Database — Aurora PostgreSQL", level=2)
    bullets(doc, [
        "Aurora PostgreSQL-Compatible, Serverless v2.",
        "Capacity: min 0.5 ACU, max 2 ACU — more than enough; keeps idle cost low.",
        "Storage auto-scales from ~10 GB; actual data footprint is well under 1 GB.",
        "Multi-AZ for resilience; automated backups with point-in-time recovery (retain 14–35 days).",
        "Note: at this volume a single db.t4g.medium standard RDS PostgreSQL instance would be cheaper than Aurora — Aurora is fine if standardized on; flagging the cost trade-off.",
    ])

    doc.add_heading("4.3  Object Storage — S3", level=2)
    bullets(doc, [
        "One bucket for documents + reports. ~40 FAIRs × 30 MB ≈ 1.2 GB; budget 50 GB for years of retention.",
        "Encryption at rest: SSE-KMS (customer-managed key).",
        "Versioning ON (protects against accidental overwrite/delete).",
        "Block all public access; access only via the application's IAM role.",
        "Lifecycle policy per Skyryse records-retention requirements.",
    ])

    doc.add_heading("4.4  Identity & Access", level=2)
    bullets(doc, [
        "Authentication must be added — the pilot has none. Preferred: integrate with Skyryse corporate SSO (SAML/OIDC) so engineers use existing company logins; alternative: AWS Cognito with MFA.",
        "Per-user identity flows into the audit log (currently records the OS user; would record the authenticated engineer).",
        "Authorize only the Supplier Quality group.",
    ])

    doc.add_heading("4.5  AI Inference — Amazon Bedrock", level=2)
    bullets(doc, [
        "Use Claude on Amazon Bedrock in the deployment region (model IDs are anthropic.-prefixed on Bedrock).",
        "The app calls Bedrock over the AWS network — controlled data does not traverse the public internet or a third-party endpoint.",
        "Confirm the required Claude model is available in the chosen GovCloud region; if not, compliance must approve an alternative.",
    ])

    doc.add_heading("4.6  Networking & Security", level=2)
    bullets(doc, [
        "VPC with private subnets for app + database; database not publicly routable.",
        "Internal Application Load Balancer with TLS (ACM certificate); access gated by corporate VPN.",
        "Security groups: ALB → app, app → Aurora (5432), app → S3/Bedrock via VPC endpoints.",
        "Secrets Manager for database credentials and any keys (no secrets in code or env files).",
        "CloudWatch logs + CloudTrail for the AWS-side audit record.",
    ])

    doc.add_heading("4.7  Backup & Disaster Recovery", level=2)
    bullets(doc, [
        "Aurora automated backups + point-in-time recovery.",
        "S3 versioning (and optional cross-region replication within GovCloud).",
        "Application audit log persisted in Aurora.",
    ])

    doc.add_heading("5. Sizing Summary", level=1)
    add_table(doc, ["Resource", "Recommended", "Rationale"], [
        ["Compute", "Fargate 2 vCPU / 4 GB ×2 (HA)", "Low concurrency; brief render spikes"],
        ["Database", "Aurora PostgreSQL Serverless v2, 0.5–2 ACU, Multi-AZ", "Tiny structured dataset; resilience"],
        ["Object storage", "S3, ~50 GB, SSE-KMS, versioned", "Documents + reports"],
        ["Identity", "Corporate SSO (SAML/OIDC) or Cognito + MFA", "3–4 named users"],
        ["Inference", "Claude via Amazon Bedrock (GovCloud)", "Keep data in AWS boundary"],
        ["Region", "AWS GovCloud (US)", "Export-controlled data"],
        ["Network", "Private VPC, internal ALB + TLS, VPN access", "Internal-only"],
    ], widths=[1.5, 3.2, 2.1])

    doc.add_heading("6. Application Changes Required (so IT can plan effort)", level=1)
    doc.add_paragraph(
        "The pilot is a working FastAPI app, but moving to this architecture requires development "
        "work — it is not a lift-and-shift:")
    bullets(doc, [
        ("Storage layer:", "replace local flat-file storage with S3 (documents/reports) + Aurora (metadata, review JSON, audit log). The storage code is isolated, so this is contained."),
        ("Authentication:", "add SSO/Cognito login and per-user identity."),
        ("Inference:", "switch the Anthropic client to the Bedrock client (anthropic[bedrock]), with anthropic.-prefixed model IDs."),
        ("Secrets:", "move DB credentials/keys to Secrets Manager."),
        ("Packaging:", "containerize (Dockerfile) for Fargate; add CloudWatch logging."),
    ], numbered=True)
    doc.add_paragraph(
        "Estimated effort for items 1–5: on the order of a few engineering days, plus IT's "
        "infrastructure provisioning. Application changes can proceed once IT provisions the AWS "
        "resources and provides connection details.")

    doc.add_heading("7. Rough Cost Ballpark", level=1)
    doc.add_paragraph(
        "For this load the AWS footprint is modest — on the order of a few hundred dollars per month "
        "(Fargate + Aurora Serverless v2 minimum + small S3 + Bedrock per-use inference). GovCloud "
        "pricing runs higher than commercial; IT should price it in the target GovCloud region. "
        "Bedrock inference is per-token and small at this volume (each review is well under a few dollars).")

    doc.add_heading("8. Open Items for IT / Compliance", level=1)
    bullets(doc, [
        "Confirm GovCloud as the deployment boundary.",
        "Confirm Claude-on-Bedrock availability in the chosen GovCloud region (vs. an approved alternative).",
        "Provide the SSO / identity integration method.",
        "Confirm records-retention period for documents, reviews, and audit logs.",
        "Confirm VPN / network access path for the engineers.",
    ])

    out = ROOT / "FAIR Review - System Requirements.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Document 2 — Compliance Summary
# ---------------------------------------------------------------------------

def build_compliance_summary():
    doc = base_doc()
    add_title(doc, "FAIR Review Tool — Data Handling Summary",
              "For Export-Control / Compliance Review")
    meta_block(doc, [
        ("Prepared for:", "Skyryse Export-Control / Compliance Officer"),
        ("Prepared by:", "Abel Napoleon, Supplier Quality"),
        ("Date:", "June 12, 2026"),
        ("Decision requested:", "Authorization to use this tool on controlled supplier data, and confirmation of permitted storage/transmission boundaries."),
    ])

    doc.add_heading("1. What the Tool Is", level=1)
    doc.add_paragraph(
        "An internal tool that helps Supplier Quality review supplier AS9102 First Article Inspection "
        "Report (FAIR) packages. A quality engineer uploads a supplier's documentation (FAIR forms, "
        "Certificates of Conformance, material certs, NADCAP certs, inspection reports, drawings), and "
        "the tool produces a draft completeness/traceability review. A human engineer makes the final "
        "acceptance decision — the AI output is decision-support, not the decision.")
    doc.add_paragraph(
        "Current deployment is a single-machine pilot (one workstation, single user).")

    doc.add_heading("2. What Data It Handles", level=1)
    doc.add_paragraph(
        "The uploaded packages contain supplier quality records and, in at least one reviewed case "
        "(part 1101021), an engineering drawing marked:")
    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Inches(0.4)
    rq = q.add_run("EAR / ECCN 7E994 — “subject to the Export Administration Regulations… export, "
                   "reexport, transfer… without proper U.S. Government authorization is strictly prohibited.”")
    rq.italic = True
    doc.add_paragraph(
        "So the tool handles EAR-controlled technical data. This summary assumes that classification governs.")

    doc.add_heading("3. Where Data Is Stored", level=1)
    add_table(doc, ["Item", "Location", "Notes"], [
        ["Original supplier documents", "Local workstation disk, parts\\<part>\\uploads\\", "Never copied elsewhere on disk"],
        ["AI review output + reports", "Same local folder", "JSON + markdown"],
        ["Daily backups", "Local FAIR_Backups\\ (ZIP)", "Can be redirected to an approved location"],
        ["Audit log", "Local audit.log", "Who/when for each action"],
    ], widths=[2.0, 2.8, 2.0])
    doc.add_paragraph(
        "No consumer cloud storage is used. Nothing leaves the workstation except as described in Section 4.")

    doc.add_heading("4. What Data Is Transmitted Externally — the Item Needing Review", level=1)
    doc.add_paragraph(
        "To generate each review, the tool transmits the contents of the uploaded documents (extracted "
        "text, and page images for scanned documents) to the Anthropic Claude API (Anthropic, a U.S.-based "
        "company) over an encrypted (HTTPS/TLS) connection. Anthropic's systems process the documents and "
        "return the review text. No other external service receives any data.")
    doc.add_paragraph("Relevant vendor facts to verify directly with Anthropic during procurement:")
    bullets(doc, [
        "Under Anthropic's commercial API terms, customer API inputs/outputs are not used to train models by default.",
        "Data retention windows and a Zero Data Retention (ZDR) option are available under enterprise agreements — the exact terms, data residency, and personnel-access controls should be confirmed in writing.",
    ])
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run("the planned multi-user AWS deployment would move inference to Amazon Bedrock inside AWS "
              "GovCloud, keeping document contents within the AWS controlled boundary rather than the public API "
              "(see the System Requirements document).")

    doc.add_heading("5. Security Controls Already in Place (pilot)", level=1)
    bullets(doc, [
        "Application bound to localhost only — not reachable from the network (verified).",
        "Audit trail of every create/upload/review action (timestamp, Windows user, action, part).",
        "Automated daily backups with rotation.",
        "Disk encryption (BitLocker) — to be enabled/confirmed on the workstation.",
        "No application login yet (single-user pilot); access is gated by the Windows account.",
    ])

    doc.add_heading("6. Decisions Requested from Compliance", level=1)
    bullets(doc, [
        ("Is transmitting ECCN 7E994 technical data to the AI service permitted", "under our EAR obligations? If conditions apply (e.g., a Zero-Data-Retention agreement, U.S.-person-only access, a signed DPA / technology-control plan), please specify them. (The planned deployment uses Amazon Bedrock in GovCloud to keep data in the AWS boundary.)"),
        ("Approved storage location", "for the controlled data going forward — local workstation only, an internal company server/share, or a compliant cloud (AWS GovCloud / Azure Government / M365 GCC High). Standard consumer/commercial cloud is assumed not acceptable."),
        ("Handling, marking, retention, or access-logging requirements", "we must apply to the stored review records and backups."),
        ("Approval scope", "— pilot (single user) now, and what is required before extending access to additional Supplier Quality engineers."),
    ], numbered=True)

    doc.add_heading("7. Contact", level=1)
    doc.add_paragraph("Abel Napoleon, Supplier Quality — abelnapoleon1@gmail.com")
    doc.add_paragraph("Tool documentation and source are retained internally with this summary.")

    out = ROOT / "FAIR Review - Compliance Summary.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print("Created:", build_system_requirements())
    print("Created:", build_compliance_summary())
