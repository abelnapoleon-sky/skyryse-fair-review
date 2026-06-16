"""Document conversion + AI FAIR review pipeline.

Converts supplier documents (PDF / Excel / Word / images) into Claude-readable
content blocks, then runs a single structured review against the Skyryse
Supplier Quality Clauses and AS9102 requirements.
"""
from __future__ import annotations

import base64
import io
import json
from pathlib import Path

import anthropic
import pandas as pd
import pdfplumber
import pypdfium2 as pdfium

MODEL = "claude-opus-4-8"
APP_DIR = Path(__file__).resolve().parent
ROOT_DIR = APP_DIR.parent
CLAUSES_PATH = ROOT_DIR / "Skyryse Quality Clauses.md"

# A scanned page typically extracts as empty or garbage; below this many
# characters we fall back to rendering the page as an image.
MIN_TEXT_CHARS_PER_PAGE = 40
RENDER_SCALE = 2.0  # ~144 DPI, enough for cert fine print


# ---------------------------------------------------------------------------
# Conversion
# ---------------------------------------------------------------------------

def _png_block(pil_image) -> dict:
    buf = io.BytesIO()
    pil_image.save(buf, format="PNG")
    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": "image/png",
            "data": base64.standard_b64encode(buf.getvalue()).decode(),
        },
    }


def convert_pdf(path: Path) -> list[dict]:
    """Text pages become text blocks; scanned pages become image blocks."""
    blocks: list[dict] = []
    with pdfplumber.open(path) as pdf:
        texts = [page.extract_text() or "" for page in pdf.pages]
    doc = pdfium.PdfDocument(str(path))
    try:
        for i, text in enumerate(texts):
            if len(text.strip()) >= MIN_TEXT_CHARS_PER_PAGE:
                blocks.append({"type": "text", "text": f"[Page {i + 1}]\n{text}"})
            else:
                blocks.append({"type": "text", "text": f"[Page {i + 1} — scanned image follows]"})
                blocks.append(_png_block(doc[i].render(scale=RENDER_SCALE).to_pil()))
    finally:
        doc.close()
    return blocks


def convert_excel(path: Path) -> list[dict]:
    sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
    parts = []
    for name, df in sheets.items():
        df = df.fillna("")
        lines = [f"## Sheet: {name}"]
        for idx, row in df.iterrows():
            cells = [str(c).strip() for c in row.tolist()]
            if any(cells):
                rendered = " | ".join(
                    f"[col {i + 1}] {c}" for i, c in enumerate(cells) if c
                )
                lines.append(f"Row {idx + 1}: {rendered}")
        parts.append("\n".join(lines))
    return [{"type": "text", "text": "\n\n".join(parts)}]


def convert_docx(path: Path) -> list[dict]:
    import docx

    d = docx.Document(str(path))
    lines = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return [{"type": "text", "text": "\n".join(lines)}]


def convert_image(path: Path) -> list[dict]:
    from PIL import Image

    return [_png_block(Image.open(path).convert("RGB"))]


def convert_file(path: Path) -> list[dict]:
    """Return Claude content blocks for one uploaded file, prefixed with a header."""
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            blocks = convert_pdf(path)
        elif ext in (".xls", ".xlsx", ".xlsm"):
            blocks = convert_excel(path)
        elif ext == ".docx":
            blocks = convert_docx(path)
        elif ext in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".tif", ".tiff"):
            blocks = convert_image(path)
        else:
            blocks = [{"type": "text", "text": path.read_text(encoding="utf-8", errors="replace")}]
    except Exception as exc:  # surface conversion failure to the reviewer
        blocks = [{"type": "text", "text": f"[CONVERSION FAILED for this file: {exc}]"}]
    header = {"type": "text", "text": f"\n===== DOCUMENT: {path.name} ====="}
    return [header, *blocks]


# ---------------------------------------------------------------------------
# Review
# ---------------------------------------------------------------------------

REVIEW_SCHEMA = {
    "type": "object",
    "properties": {
        "verdict": {"type": "string", "enum": ["APPROVED", "UNAPPROVED", "NEEDS_INFO"]},
        "part_number": {"type": "string"},
        "part_name": {"type": "string"},
        "supplier": {"type": "string"},
        "po_number": {"type": "string"},
        "summary": {
            "type": "string",
            "description": "Short plain-English summary: does the part have full traceability from the OEM to Skyryse? If not, why not.",
        },
        "traceability_complete": {"type": "boolean"},
        "traceability_chain": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "step": {"type": "string", "description": "Role in the chain, e.g. Mill, Distributor, OCM, Special Process"},
                    "entity": {"type": "string"},
                    "document": {"type": "string", "description": "Document name and page that evidences this step"},
                    "key_ids": {"type": "string", "description": "Heat/lot/cert/PO numbers tying this step to the next"},
                    "linked": {"type": "boolean", "description": "Is this step documentarily linked to the next step?"},
                },
                "required": ["step", "entity", "document", "key_ids", "linked"],
                "additionalProperties": False,
            },
        },
        "findings": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "severity": {"type": "string", "enum": ["major", "minor", "observation"]},
                    "category": {"type": "string", "description": "e.g. Traceability, AS9102 Form 1/2/3, Special Process, CoC, Material"},
                    "description": {"type": "string"},
                    "location": {"type": "string", "description": "Exact document name and page/sheet/block where the reviewer can verify this"},
                    "required_action": {"type": "string"},
                },
                "required": ["severity", "category", "description", "location", "required_action"],
                "additionalProperties": False,
            },
        },
        "clarifying_questions": {"type": "array", "items": {"type": "string"}},
    },
    "required": [
        "verdict", "part_number", "part_name", "supplier", "po_number", "summary",
        "traceability_complete", "traceability_chain", "findings", "clarifying_questions",
    ],
    "additionalProperties": False,
}


def _system_prompt() -> str:
    clauses = ""
    if CLAUSES_PATH.exists():
        clauses = CLAUSES_PATH.read_text(encoding="utf-8")
    return f"""You are a Quality Engineer at Skyryse experienced in AS9102 First Article Inspection Report (FAIR) review and acceptance. You review FAIR packages submitted by suppliers for completeness, accuracy, continuity, and traceability.

Review rules:
- Treat the Skyryse Supplier Quality Clauses below as flow-down requirements, alongside AS9102 (latest revision) and all controlling drawings, specifications, and purchase order requirements that appear in the package.
- Verify FULL traceability from the OEM/mill through every intermediary (every distributor must have a transactional document — packing slip, invoice, or cert — linking it to the next party) to Skyryse as the customer. A missing intermediary link is a major finding.
- Verify AS9102 Forms 1, 2, and 3 are complete and signed: all characteristics have ACTUAL recorded results (not restated requirements), pass/fail dispositions are indicated, special processes appear on Form 2 with their certificate numbers, serial/lot numbers are assigned where required, and any nonconformance carries an approved deviation (SEP) ticket number.
- Verify NADCAP claims: any accreditation scope cited on the forms must actually appear on the provided NADCAP certificate, and the certificate must not be expired relative to the work performed.
- Verify CoC minimum content (supplier, PO, part number/rev, lot/serial, quantity, specs, conformity statement, signature, date).
- Cross-check heat numbers, lot numbers, cert numbers, quantities, and dates across ALL documents; flag inconsistencies.
- For every finding, give the EXACT location (document name + page/sheet/form block) so a human reviewer can verify it.
- Mark the verdict UNAPPROVED if documentation is missing, incomplete, illegible, or inconsistent; if entries do not comply with drawings/specs/PO/revisions; or if deviations exist without prior Skyryse approval. Use NEEDS_INFO only when you cannot reach a verdict without an answer from the Skyryse engineer (not the supplier).
- Be precise and conservative: never invent document content. If a page is illegible, say so and list it as a finding.

SKYRYSE SUPPLIER QUALITY CLAUSES (flow-down requirements):
{clauses}
"""


def run_review(part_dir: Path, meta: dict) -> dict:
    """Convert all uploaded files for a part and run the Claude review."""
    uploads = sorted((part_dir / "uploads").glob("*"))
    if not uploads:
        raise ValueError("No files uploaded for this part")

    content: list[dict] = [{
        "type": "text",
        "text": (
            "Review the following FAIR documentation package for one part.\n"
            f"Part record created by the Skyryse reviewer: part number: {meta.get('part_number', 'unknown')}, "
            f"revision: {meta.get('revision', 'unknown')}, supplier: {meta.get('supplier', 'unknown')}, "
            f"PO: {meta.get('po_number', 'unknown')}.\n"
            f"Reviewer notes: {meta.get('notes') or 'none'}\n"
            "Scanned pages are provided as images — read them carefully, including handwritten "
            "annotations, stamps, and signatures. Some scans may be rotated."
        ),
    }]
    for f in uploads:
        content.extend(convert_file(f))

    client = anthropic.Anthropic()
    with client.messages.stream(
        model=MODEL,
        max_tokens=32000,
        thinking={"type": "adaptive"},
        output_config={
            "effort": "high",
            "format": {"type": "json_schema", "schema": REVIEW_SCHEMA},
        },
        system=[{
            "type": "text",
            "text": _system_prompt(),
            "cache_control": {"type": "ephemeral"},
        }],
        messages=[{"role": "user", "content": content}],
    ) as stream:
        message = stream.get_final_message()

    text = next(b.text for b in message.content if b.type == "text")
    review = json.loads(text)
    review["model"] = message.model
    review["usage"] = {
        "input_tokens": message.usage.input_tokens,
        "output_tokens": message.usage.output_tokens,
        "cache_read_input_tokens": message.usage.cache_read_input_tokens,
        "cache_creation_input_tokens": message.usage.cache_creation_input_tokens,
    }
    return review


METADATA_SCHEMA = {
    "type": "object",
    "properties": {
        "part_number": {"type": "string", "description": "Part number exactly as printed; empty string if not found anywhere in the package"},
        "part_name": {"type": "string", "description": "Part name / nomenclature / description from the drawing or FAIR Form 1; empty if not found"},
        "revision": {"type": "string", "description": "Part/drawing revision level; empty if not found"},
        "supplier": {"type": "string", "description": "Manufacturing supplier (the company that made/processed the part and signs the CoC); empty if not found"},
        "po_number": {"type": "string", "description": "Skyryse Purchase Order number; empty if not found"},
        "found_part_number": {"type": "boolean", "description": "True if a part number appears in the documents"},
        "found_supplier": {"type": "boolean"},
        "found_po_number": {"type": "boolean"},
        "discrepancy_note": {"type": "string", "description": "Note any inconsistency between documents for these fields, e.g. PO differs between CoC and packing slip; empty if consistent"},
    },
    "required": [
        "part_number", "part_name", "revision", "supplier", "po_number",
        "found_part_number", "found_supplier", "found_po_number", "discrepancy_note",
    ],
    "additionalProperties": False,
}


def extract_metadata(part_dir: Path) -> dict:
    """Read the package and pull identifying fields exactly as printed in the docs."""
    uploads = sorted((part_dir / "uploads").glob("*"))
    if not uploads:
        raise ValueError("No files uploaded for this part")

    content: list[dict] = [{
        "type": "text",
        "text": (
            "Read this supplier AS9102 FAIR documentation package and extract the identifying "
            "fields EXACTLY as they appear in the documents. Use the controlling drawing, FAIR "
            "Form 1, and the Certificate of Conformance as the authority. For part_name use the "
            "nomenclature/description on the drawing or Form 1. The supplier is the company that "
            "manufactured/processed the part and signs the CoC (not a distributor or the mill). "
            "If a field does not appear anywhere, return an empty string and set its found_* flag "
            "to false. Note in discrepancy_note if the same field shows different values across "
            "documents. Scanned pages are images — read them, including stamps and handwriting."
        ),
    }]
    for f in uploads:
        content.extend(convert_file(f))

    client = anthropic.Anthropic()
    msg = client.messages.create(
        model=MODEL,
        max_tokens=2000,
        output_config={
            "effort": "low",
            "format": {"type": "json_schema", "schema": METADATA_SCHEMA},
        },
        system="You extract identifying metadata from AS9102 FAIR packages. Report values exactly as printed in the documents; never guess.",
        messages=[{"role": "user", "content": content}],
    )
    text = next(b.text for b in msg.content if b.type == "text")
    return json.loads(text)


def review_to_markdown(review: dict, meta: dict) -> str:
    """Render the structured review as a human-readable markdown report."""
    lines = [
        f"# FAIR Review — {review.get('part_number')} — {review.get('part_name')}",
        "",
        f"**Verdict:** {review.get('verdict')}  |  **Supplier:** {review.get('supplier')}  |  **PO:** {review.get('po_number')}",
        "",
        "## Summary",
        review.get("summary", ""),
        "",
        f"## Traceability chain (complete: {'YES' if review.get('traceability_complete') else 'NO'})",
        "",
        "| Step | Entity | Document | Key IDs | Linked |",
        "|---|---|---|---|---|",
    ]
    for s in review.get("traceability_chain", []):
        lines.append(
            f"| {s['step']} | {s['entity']} | {s['document']} | {s['key_ids']} | {'✓' if s['linked'] else '✗ BROKEN'} |"
        )
    lines += ["", "## Findings", ""]
    for i, f in enumerate(review.get("findings", []), 1):
        lines += [
            f"### {i}. [{f['severity'].upper()}] {f['category']}",
            f"- **Issue:** {f['description']}",
            f"- **Where to verify:** {f['location']}",
            f"- **Required action:** {f['required_action']}",
            "",
        ]
    questions = review.get("clarifying_questions", [])
    if questions:
        lines += ["## Clarifying questions", ""]
        lines += [f"- {q}" for q in questions]
    return "\n".join(lines)
